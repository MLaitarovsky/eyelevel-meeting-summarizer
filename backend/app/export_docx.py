from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from app.schemas import MeetingAnalysis

# Hebrew block: U+0590..U+05FF. We treat any Hebrew presence as RTL — fine for
# pure-Hebrew and Hebrew-dominant lines; the rare mixed-language paragraph still
# reads correctly when right-aligned.
_HEBREW_RE = re.compile("[֐-׿]")


def _is_rtl(text: str) -> bool:
    return bool(_HEBREW_RE.search(text))


def _add_paragraph(doc, text: str, *, style: str | None = None):
    p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    if _is_rtl(text):
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def to_docx(analysis: MeetingAnalysis, transcript: str) -> bytes:
    doc = Document()

    title = doc.add_heading(analysis.title, level=1)
    if _is_rtl(analysis.title):
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading("Summary", level=2)
    for chunk in re.split(r"\n\s*\n", analysis.summary):
        chunk = chunk.strip()
        if chunk:
            _add_paragraph(doc, chunk)

    doc.add_heading("Participants", level=2)
    if analysis.participants:
        for p in analysis.participants:
            line = f"{p.name} — {p.role}" if p.role else p.name
            _add_paragraph(doc, line, style="List Bullet")
    else:
        doc.add_paragraph("(none identified)")

    doc.add_heading("Decisions", level=2)
    if analysis.decisions:
        for d in analysis.decisions:
            _add_paragraph(doc, d.decision, style="List Number")
            if d.context:
                ctx = doc.add_paragraph(d.context)
                ctx.paragraph_format.left_indent = Inches(0.5)
                if _is_rtl(d.context):
                    ctx.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        doc.add_paragraph("(none)")

    doc.add_heading("Action Items", level=2)
    if analysis.action_items:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid"
        header = table.rows[0].cells
        header[0].text = "Task"
        header[1].text = "Owner"
        header[2].text = "Deadline"
        header[3].text = "Priority"
        for item in analysis.action_items:
            row = table.add_row().cells
            row[0].text = item.task
            row[1].text = item.owner or ""
            row[2].text = item.deadline or ""
            row[3].text = item.priority or ""
    else:
        doc.add_paragraph("(none)")

    doc.add_heading("Open Questions", level=2)
    if analysis.open_questions:
        for q in analysis.open_questions:
            _add_paragraph(doc, q, style="List Bullet")
    else:
        doc.add_paragraph("(none)")

    doc.add_page_break()
    doc.add_heading("Full Transcript", level=2)
    for line in transcript.splitlines():
        line = line.strip()
        if line:
            _add_paragraph(doc, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
